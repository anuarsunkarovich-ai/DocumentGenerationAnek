"""Live backend smoke tests against a running Docker stack."""

from __future__ import annotations

import os
import time
from io import BytesIO
from uuid import uuid4
from zipfile import ZipFile

import httpx
import pytest
from docx import Document

pytestmark = [pytest.mark.integration, pytest.mark.live_stack]


def build_docx_fixture(*document_text: str) -> bytes:
    """Build a minimal DOCX fixture for live upload and generation checks."""
    body = "".join(f"<w:t>{text}</w:t>" for text in document_text)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}</w:body>"
        "</w:document>"
    )
    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types></Types>")
        archive.writestr("word/document.xml", document_xml)
    return buffer.getvalue()


def build_assisted_templateization_docx() -> tuple[bytes, dict[str, str], dict[str, dict[str, object]]]:
    """Build a narrative DOCX fixture plus deterministic selection spans."""
    document = Document()
    document.add_paragraph("Отчет по учебной практике")
    document.add_paragraph("Кафедра информационных систем")
    document.add_paragraph("Тема: Экономический анализ цифровых расходов")
    document.add_paragraph("Выполнил: Иванов Иван студент группы ИС-101")
    document.add_paragraph("Научный руководитель: Петров А.А.")
    document.add_paragraph("Алматы 2026г")

    buffer = BytesIO()
    document.save(buffer)

    paragraphs = {
        "topic": "Тема: Экономический анализ цифровых расходов",
        "student": "Выполнил: Иванов Иван студент группы ИС-101",
        "supervisor": "Научный руководитель: Петров А.А.",
        "city_year": "Алматы 2026г",
    }
    replacements = {
        "report_topic": "Влияние цифровых сервисов на учет расходов",
        "student_identity": "Сериков Нұржан студент группы ИС-404",
        "supervisor_name": "Ахметова Л.С.",
        "city_year": "Кызылорда 2027г",
    }
    selections = {
        "report_topic": {
            "paragraph_path": "body/p/2",
            "fragment_start": paragraphs["topic"].index("Экономический"),
            "fragment_end": len(paragraphs["topic"]),
            "label": "Report Topic",
        },
        "student_identity": {
            "paragraph_path": "body/p/3",
            "fragment_start": paragraphs["student"].index("Иванов"),
            "fragment_end": len(paragraphs["student"]),
            "label": "Student Identity",
        },
        "supervisor_name": {
            "paragraph_path": "body/p/4",
            "fragment_start": paragraphs["supervisor"].index("Петров"),
            "fragment_end": len(paragraphs["supervisor"]),
            "label": "Supervisor Name",
        },
        "city_year": {
            "paragraph_path": "body/p/5",
            "fragment_start": 0,
            "fragment_end": len(paragraphs["city_year"]),
            "label": "City And Year",
        },
    }
    return buffer.getvalue(), replacements, selections


@pytest.fixture(scope="module")
def live_stack_config() -> dict[str, str]:
    """Return the environment required to run live-stack smoke tests."""
    if os.getenv("RUN_LIVE_STACK") != "1":
        pytest.skip("Set RUN_LIVE_STACK=1 to run live stack smoke tests.")

    required = {
        "api_base_url": os.getenv("LIVE_API_BASE_URL", "http://127.0.0.1:8000"),
        "email": os.getenv("LIVE_ADMIN_EMAIL"),
        "password": os.getenv("LIVE_ADMIN_PASSWORD"),
        "organization_id": os.getenv("LIVE_ORGANIZATION_ID"),
    }
    missing = [key for key, value in required.items() if not value]
    if missing:
        pytest.fail(f"Missing live stack environment values: {', '.join(sorted(missing))}")
    return {key: value for key, value in required.items() if value is not None}


def wait_for_job_completion(
    client: httpx.Client,
    *,
    task_id: str,
    organization_id: str,
    headers: dict[str, str],
    timeout_seconds: float = 30.0,
) -> dict:
    """Poll one generation job until it reaches a terminal state."""
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        response = client.get(
            f"/api/v1/documents/jobs/{task_id}",
            params={"organization_id": organization_id},
            headers=headers,
        )
        response.raise_for_status()
        payload = response.json()
        if payload["status"] in {"completed", "failed"}:
            return payload
        time.sleep(0.5)
    pytest.fail(f"Timed out waiting for job {task_id} to complete.")


def test_live_stack_end_to_end_flow(live_stack_config: dict[str, str]) -> None:
    """Verify the live Docker stack from health checks through cached generation."""
    unique_suffix = uuid4().hex[:8]
    template_code = f"smoke-certificate-{unique_suffix}"
    template_name = f"Smoke Certificate {unique_suffix}"
    template_version = "1.0.0"
    constructor = {
        "locale": "en-US",
        "metadata": {"document_type": "certificate"},
        "blocks": [
            {
                "type": "header",
                "id": "header-1",
                "text": "Certificate",
            },
            {
                "type": "text",
                "id": "text-1",
                "binding": {"key": "student_name"},
            },
        ],
    }
    data_payload = {
        "student_name": "Anek",
    }

    with httpx.Client(base_url=live_stack_config["api_base_url"], timeout=30.0) as client:
        live_response = client.get("/health/live")
        assert live_response.status_code == 200
        assert live_response.json()["status"] == "ok"

        ready_response = client.get("/health/ready")
        assert ready_response.status_code == 200
        ready_payload = ready_response.json()
        assert ready_payload["checks"]["database"]["status"] == "ok"
        assert ready_payload["checks"]["storage"]["status"] == "ok"
        assert ready_payload["checks"]["redis"]["status"] == "ok"

        login_response = client.post(
            "/api/v1/auth/login",
            json={
                "email": live_stack_config["email"],
                "password": live_stack_config["password"],
            },
        )
        assert login_response.status_code == 200
        login_payload = login_response.json()
        access_token = login_payload["access_token"]
        refresh_token = login_payload["refresh_token"]
        headers = {"Authorization": f"Bearer {access_token}"}

        schema_response = client.post(
            "/api/v1/templates/extract-schema",
            headers=headers,
            files={
                "file": (
                    "certificate.docx",
                    build_docx_fixture("{{student_name}}"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert schema_response.status_code == 200
        schema_payload = schema_response.json()
        assert schema_payload["variable_count"] == 1

        upload_response = client.post(
            "/api/v1/templates/upload",
            headers=headers,
            data={
                "organization_id": live_stack_config["organization_id"],
                "name": template_name,
                "code": template_code,
                "version": template_version,
                "description": "Live smoke test upload",
                "publish": "true",
            },
            files={
                "file": (
                    "certificate.docx",
                    build_docx_fixture("{{student_name}}"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert upload_response.status_code == 201
        upload_payload = upload_response.json()
        template_id = upload_payload["template"]["id"]
        template_version_id = upload_payload["version"]["id"]

        stored_schema_response = client.post(
            f"/api/v1/templates/{template_id}/extract-schema",
            headers=headers,
            params={"organization_id": live_stack_config["organization_id"]},
        )
        assert stored_schema_response.status_code == 200
        assert stored_schema_response.json()["schema"]["variable_count"] == 1

        generate_response = client.post(
            "/api/v1/documents/generate",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "template_id": template_id,
                "template_version_id": template_version_id,
                "data": data_payload,
                "constructor": constructor,
            },
        )
        assert generate_response.status_code == 202
        assert generate_response.headers["X-Request-ID"]
        assert generate_response.headers["X-Correlation-ID"]
        generate_payload = generate_response.json()
        assert generate_payload["from_cache"] is False
        assert generate_payload["status"] == "queued"

        completed_job = wait_for_job_completion(
            client,
            task_id=generate_payload["task_id"],
            organization_id=live_stack_config["organization_id"],
            headers=headers,
        )
        assert completed_job["status"] == "completed"
        assert completed_job["from_cache"] is False
        assert completed_job["artifacts"]

        download_response = client.get(
            f"/api/v1/documents/jobs/{generate_payload['task_id']}/download",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert download_response.status_code == 200
        download_payload = download_response.json()
        assert download_payload["artifact"]["download_url"]
        downloaded_artifact = httpx.get(download_payload["artifact"]["download_url"], timeout=30.0)
        assert downloaded_artifact.status_code == 200
        assert len(downloaded_artifact.content) > 0

        verify_response = client.post(
            "/api/v1/documents/verify",
            headers=headers,
            data={"organization_id": live_stack_config["organization_id"]},
            files={
                "file": (
                    download_payload["artifact"]["file_name"],
                    downloaded_artifact.content,
                    download_payload["artifact"]["content_type"],
                )
            },
        )
        assert verify_response.status_code == 200
        verify_payload = verify_response.json()
        assert verify_payload["matched"] is True
        assert verify_payload["artifact"]["artifact_id"] == download_payload["artifact"]["id"]
        assert verify_payload["artifact"]["task_id"] == generate_payload["task_id"]
        assert verify_payload["artifact"]["verification_code"]
        authenticity_hash = verify_payload["artifact"]["authenticity_hash"]

        preview_response = client.get(
            f"/api/v1/documents/jobs/{generate_payload['task_id']}/preview",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert preview_response.status_code == 200
        preview_url = preview_response.json()["artifact"]["download_url"]
        preview_artifact = httpx.get(preview_url, timeout=30.0)
        assert preview_artifact.status_code == 200
        assert len(preview_artifact.content) > 0

        audit_response = client.get(
            "/api/v1/admin/diagnostics/audit-events",
            params={
                "organization_id": live_stack_config["organization_id"],
                "limit": 25,
            },
            headers=headers,
        )
        assert audit_response.status_code == 200
        audit_payload = audit_response.json()
        assert any(
            event["action"] == "document_job_completed" for event in audit_payload["items"]
        )

        api_key_response = client.post(
            "/api/v1/admin/api-keys",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "name": "Live smoke key",
                "scopes": [
                    "templates:read",
                    "documents:generate",
                    "documents:read",
                    "audit:read",
                ],
            },
        )
        assert api_key_response.status_code == 201
        api_key_payload = api_key_response.json()
        public_headers = {"X-API-Key": api_key_payload["api_key"]}

        public_templates_response = client.get(
            "/api/v1/public/templates",
            headers=public_headers,
        )
        assert public_templates_response.status_code == 200
        assert public_templates_response.json()["items"]

        public_verify_response = client.post(
            "/api/v1/public/documents/verify",
            headers=public_headers,
            data={"authenticity_hash": authenticity_hash},
        )
        assert public_verify_response.status_code == 200
        assert public_verify_response.json()["matched"] is True

        audit_history_response = client.get(
            "/api/v1/admin/support/audit-history",
            params={
                "organization_id": live_stack_config["organization_id"],
                "entity_type": "document_job",
                "entity_id": generate_payload["task_id"],
            },
            headers=headers,
        )
        assert audit_history_response.status_code == 200
        assert audit_history_response.json()["items"]

        invalidate_cache_response = client.post(
            f"/api/v1/admin/support/jobs/{generate_payload['task_id']}/invalidate-cache",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert invalidate_cache_response.status_code == 200
        assert invalidate_cache_response.json()["invalidated_artifact_count"] >= 1

        cache_before_response = client.get(
            "/api/v1/admin/diagnostics/cache-stats",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert cache_before_response.status_code == 200
        assert cache_before_response.json()["completed_jobs"] >= 1

        time.sleep(1.0)

        regenerated_response = client.post(
            "/api/v1/documents/generate",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "template_id": template_id,
                "template_version_id": template_version_id,
                "data": data_payload,
                "constructor": constructor,
            },
        )
        assert regenerated_response.status_code == 202
        assert regenerated_response.json()["from_cache"] is False
        regenerated_job = wait_for_job_completion(
            client,
            task_id=regenerated_response.json()["task_id"],
            organization_id=live_stack_config["organization_id"],
            headers=headers,
        )
        assert regenerated_job["status"] == "completed"

        cached_generate_response = client.post(
            "/api/v1/documents/generate",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "template_id": template_id,
                "template_version_id": template_version_id,
                "data": data_payload,
                "constructor": constructor,
            },
        )
        assert cached_generate_response.status_code == 202
        cached_payload = cached_generate_response.json()
        assert cached_payload["from_cache"] is True
        assert cached_payload["status"] == "completed"

        cleanup_response = client.post(
            "/api/v1/admin/support/maintenance/cleanup",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert cleanup_response.status_code == 200

        cache_after_response = client.get(
            "/api/v1/admin/diagnostics/cache-stats",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert cache_after_response.status_code == 200
        cache_after_payload = cache_after_response.json()
        assert cache_after_payload["cached_jobs"] >= 1
        assert cache_after_payload["cached_artifacts"] >= 1

        worker_response = client.get(
            "/api/v1/admin/diagnostics/worker-status",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert worker_response.status_code == 200
        worker_payload = worker_response.json()
        assert worker_payload["queue_depth"] >= 0
        assert worker_payload["workers"]
        assert any(worker["is_online"] for worker in worker_payload["workers"])

        metrics_response = client.get("/metrics")
        assert metrics_response.status_code == 200
        assert "document_job_result_total" in metrics_response.text
        assert "http_request_latency_seconds" in metrics_response.text

        refresh_response = client.post(
            "/api/v1/auth/refresh",
            json={"refresh_token": refresh_token},
        )
        assert refresh_response.status_code == 200
        refreshed_payload = refresh_response.json()

        logout_response = client.post(
            "/api/v1/auth/logout",
            headers={"Authorization": f"Bearer {refreshed_payload['access_token']}"},
            json={"refresh_token": refreshed_payload["refresh_token"]},
        )
        assert logout_response.status_code == 204


def test_live_stack_assisted_templateization_flow(live_stack_config: dict[str, str]) -> None:
    """Verify the live stack can templateize a normal DOCX and generate from saved spans."""
    source_docx, replacements, selections = build_assisted_templateization_docx()
    unique_suffix = uuid4().hex[:8]
    template_code = f"smoke-import-{unique_suffix}"
    template_name = f"Smoke Import {unique_suffix}"

    with httpx.Client(base_url=live_stack_config["api_base_url"], timeout=30.0) as client:
        login_response = client.post(
            "/api/v1/auth/login",
            json={
                "email": live_stack_config["email"],
                "password": live_stack_config["password"],
            },
        )
        assert login_response.status_code == 200
        access_token = login_response.json()["access_token"]
        headers = {"Authorization": f"Bearer {access_token}"}

        inspect_upload_response = client.post(
            "/api/v1/templates/import/inspect",
            headers=headers,
            files={
                "file": (
                    "report.docx",
                    source_docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert inspect_upload_response.status_code == 200
        inspect_upload_payload = inspect_upload_response.json()
        assert inspect_upload_payload["paragraph_count"] == 6

        analyze_upload_response = client.post(
            "/api/v1/templates/import/analyze",
            headers=headers,
            files={
                "file": (
                    "report.docx",
                    source_docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert analyze_upload_response.status_code == 200
        assert analyze_upload_response.json()["candidate_count"] == 0

        upload_response = client.post(
            "/api/v1/templates/upload",
            headers=headers,
            data={
                "organization_id": live_stack_config["organization_id"],
                "name": template_name,
                "code": template_code,
                "version": "1.0.0",
                "description": "Live assisted templateization smoke test",
                "publish": "true",
            },
            files={
                "file": (
                    "report.docx",
                    source_docx,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert upload_response.status_code == 201
        upload_payload = upload_response.json()
        template_id = upload_payload["template"]["id"]
        template_version_id = upload_payload["version"]["id"]

        inspect_stored_response = client.post(
            f"/api/v1/templates/{template_id}/import/inspect",
            headers=headers,
            json={"organization_id": live_stack_config["organization_id"]},
        )
        assert inspect_stored_response.status_code == 200
        inspect_stored_payload = inspect_stored_response.json()
        assert inspect_stored_payload["inspection_checksum"] == inspect_upload_payload["inspection_checksum"]
        assert inspect_stored_payload["paragraph_count"] == 6

        templateize_response = client.post(
            f"/api/v1/templates/{template_id}/import/templateize",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "inspection_checksum": inspect_stored_payload["inspection_checksum"],
                "selections": [
                    {
                        "paragraph_path": selection["paragraph_path"],
                        "fragment_start": selection["fragment_start"],
                        "fragment_end": selection["fragment_end"],
                        "binding_key": binding_key,
                        "label": selection["label"],
                    }
                    for binding_key, selection in selections.items()
                ],
            },
        )
        assert templateize_response.status_code == 200
        templateize_payload = templateize_response.json()
        assert templateize_payload["render_strategy"] == "docx_import"
        assert templateize_payload["confirmed_binding_count"] == 4
        assert {item["binding_key"] for item in templateize_payload["bindings"]} == set(replacements)

        template_detail_response = client.get(
            f"/api/v1/templates/{template_id}",
            headers=headers,
            params={"organization_id": live_stack_config["organization_id"]},
        )
        assert template_detail_response.status_code == 200
        current_version = template_detail_response.json()["current_version_details"]
        assert current_version["render_strategy"] == "docx_import"
        assert current_version["imported_binding_count"] == 4

        generate_response = client.post(
            "/api/v1/documents/generate-imported",
            headers=headers,
            json={
                "organization_id": live_stack_config["organization_id"],
                "template_id": template_id,
                "template_version_id": template_version_id,
                "data": replacements,
            },
        )
        assert generate_response.status_code == 202
        generate_payload = generate_response.json()
        assert generate_payload["status"] == "queued"
        assert generate_payload["from_cache"] is False

        completed_job = wait_for_job_completion(
            client,
            task_id=generate_payload["task_id"],
            organization_id=live_stack_config["organization_id"],
            headers=headers,
        )
        assert completed_job["status"] == "completed"
        assert [artifact["kind"] for artifact in completed_job["artifacts"]] == ["docx"]

        download_response = client.get(
            f"/api/v1/documents/jobs/{generate_payload['task_id']}/download",
            params={"organization_id": live_stack_config["organization_id"]},
            headers=headers,
        )
        assert download_response.status_code == 200
        download_payload = download_response.json()
        assert download_payload["artifact"]["kind"] == "docx"

        downloaded_artifact = httpx.get(download_payload["artifact"]["download_url"], timeout=30.0)
        assert downloaded_artifact.status_code == 200

        rendered = Document(BytesIO(downloaded_artifact.content))
        rendered_paragraphs = [paragraph.text for paragraph in rendered.paragraphs if paragraph.text]
        assert f"Тема: {replacements['report_topic']}" in rendered_paragraphs
        assert f"Выполнил: {replacements['student_identity']}" in rendered_paragraphs
        assert f"Научный руководитель: {replacements['supervisor_name']}" in rendered_paragraphs
        assert replacements["city_year"] in rendered_paragraphs
