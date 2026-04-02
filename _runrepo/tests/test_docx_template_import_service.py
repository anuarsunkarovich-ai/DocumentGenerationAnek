"""Tests for regular DOCX import analysis and preserved-layout rendering."""

from io import BytesIO
from uuid import uuid4

from docx import Document

from app.dtos.template import (
    TemplateImportBindingConfirmationItem,
    TemplateImportManualSelectionItem,
)
from app.services.docx_template_import_service import DocxTemplateImportService
from app.services.generation.docx_template_render_service import DocxTemplateRenderService
from app.services.generation.models import ResolvedTemplateContext


def test_docx_import_inspection_lists_addressable_paragraphs() -> None:
    """Inspection should expose plain paragraphs and table-cell paragraphs."""
    service = DocxTemplateImportService()

    inspection = service.inspect("invoice.docx", _build_import_docx())

    assert inspection.paragraph_count >= 3
    assert inspection.paragraphs[0].text == "Client Name: __________"
    assert any(paragraph.path.startswith("body/t/0/") for paragraph in inspection.paragraphs)


def test_docx_import_analysis_detects_blank_bracketed_and_table_fields() -> None:
    """The importer should detect likely fields in ordinary DOCX content."""
    service = DocxTemplateImportService()

    analysis = service.analyze("invoice.docx", _build_import_docx())

    assert analysis.candidate_count == 3
    assert {candidate.suggested_binding for candidate in analysis.candidates} == {
        "amount",
        "client_name",
        "invoice_date",
    }
    assert any(candidate.detection_kind == "blank_span" for candidate in analysis.candidates)
    assert any(candidate.detection_kind == "bracketed_label" for candidate in analysis.candidates)
    assert any(candidate.paragraph_path.startswith("body/t/0/") for candidate in analysis.candidates)


def test_docx_import_analysis_skips_citation_style_brackets_and_transliterates_cyrillic() -> None:
    """The importer should ignore citations and build stable keys from Cyrillic labels."""
    document = Document()
    document.add_paragraph("Согласно исследованию PostgreSQL эффективна [4, 11].")
    document.add_paragraph("Пользователь: [Пользователь]")

    buffer = BytesIO()
    document.save(buffer)

    service = DocxTemplateImportService()
    analysis = service.analyze("report.docx", buffer.getvalue())

    assert analysis.candidate_count == 1
    assert analysis.candidates[0].label == "Пользователь"
    assert analysis.candidates[0].suggested_binding == "polzovatel"


def test_docx_manual_templateization_builds_bindings_from_selected_spans() -> None:
    """Manual selections should turn ordinary DOCX text into confirmed bindings."""
    service = DocxTemplateImportService()
    inspection = service.inspect("invoice.docx", _build_import_docx())

    selections = [
        TemplateImportManualSelectionItem(
            paragraph_path="body/p/0",
            fragment_start=13,
            fragment_end=23,
            binding_key="client_name",
            label="Client Name",
        ),
        TemplateImportManualSelectionItem(
            paragraph_path="body/p/1",
            fragment_start=14,
            fragment_end=28,
            binding_key="invoice_date",
            label="Invoice Date",
        ),
    ]

    bindings, schema = service.templateize_from_selections(
        inspection=inspection,
        selections=selections,
    )

    assert [binding.binding_key for binding in bindings] == ["client_name", "invoice_date"]
    assert bindings[0].raw_fragment == "__________"
    assert bindings[1].raw_fragment == "[Invoice Date]"
    assert schema.variable_count == 2


def test_docx_import_confirmation_builds_bindings_and_schema() -> None:
    """Confirmed bindings should produce the persisted schema used for generation."""
    service = DocxTemplateImportService()
    analysis = service.analyze("invoice.docx", _build_import_docx())

    confirmations = [
        TemplateImportBindingConfirmationItem(
            candidate_id=candidate.id,
            binding_key=candidate.suggested_binding,
            label=candidate.label,
        )
        for candidate in analysis.candidates
    ]

    bindings, schema = service.confirm_bindings(
        analysis=analysis,
        confirmations=confirmations,
    )

    assert len(bindings) == 3
    assert schema.variable_count == 3
    assert {item.key for item in schema.variables} == {
        "amount",
        "client_name",
        "invoice_date",
    }


def test_docx_import_renderer_preserves_docx_structure_while_filling_values() -> None:
    """Rendering should replace only the detected fragments in the original DOCX."""
    source_bytes = _build_import_docx()
    import_service = DocxTemplateImportService()
    analysis = import_service.analyze("invoice.docx", source_bytes)
    confirmations = [
        TemplateImportBindingConfirmationItem(
            candidate_id=candidate.id,
            binding_key=candidate.suggested_binding,
            label=candidate.label,
        )
        for candidate in analysis.candidates
    ]
    bindings, schema = import_service.confirm_bindings(
        analysis=analysis,
        confirmations=confirmations,
    )
    context = ResolvedTemplateContext(
        template_id=uuid4(),
        template_version_id=uuid4(),
        organization_id=uuid4(),
        organization_code="finance",
        template_code="invoice",
        template_name="Invoice",
        template_version="1.0.0",
        original_filename="invoice.docx",
        variable_schema=schema.model_dump(mode="json"),
        render_strategy="docx_import",
        import_bindings=[binding.model_dump(mode="json") for binding in bindings],
    )

    renderer = DocxTemplateRenderService()
    normalized_payload, _ = renderer.prepare_payload(
        context=context,
        data={
            "client_name": "Anek LLC",
            "invoice_date": "2026-03-21",
            "amount": "1500 USD",
        },
    )
    rendered = renderer.render(
        content=source_bytes,
        context=context,
        data=normalized_payload,
    )

    document = Document(BytesIO(rendered))
    assert document.paragraphs[0].text == "Client Name: Anek LLC"
    assert document.paragraphs[1].text == "Invoice Date: 2026-03-21"
    assert document.tables[0].rows[1].cells[1].text == "1500 USD"


def _build_import_docx() -> bytes:
    """Create a realistic DOCX with ordinary fillable regions instead of placeholders."""
    document = Document()
    document.add_paragraph("Client Name: __________")
    document.add_paragraph("Invoice Date: [Invoice Date]")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Amount"
    table.rows[1].cells[0].text = "Consulting"
    table.rows[1].cells[1].text = "....."

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
