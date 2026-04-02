"""Shared helpers for traversing and locating DOCX paragraph targets."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document


@dataclass(slots=True)
class ParagraphTarget:
    """A stable paragraph reference inside a DOCX container."""

    path: str
    paragraph: Any
    source_type: str
    table_header_label: str | None = None


def load_docx_document(content: bytes):
    """Load a DOCX file into a python-docx document."""
    return Document(BytesIO(content))


def iter_document_paragraph_targets(document) -> list[ParagraphTarget]:
    """Return all addressable paragraphs in body, headers, and footers."""
    targets: list[ParagraphTarget] = []
    targets.extend(_iter_container_paragraphs(document, base_path="body", source_type="body"))
    for section_index, section in enumerate(document.sections):
        targets.extend(
            _iter_container_paragraphs(
                section.header,
                base_path=f"header/{section_index}",
                source_type="header",
            )
        )
        targets.extend(
            _iter_container_paragraphs(
                section.footer,
                base_path=f"footer/{section_index}",
                source_type="footer",
            )
        )
    return targets


def resolve_paragraph_target(document, path: str):
    """Resolve one paragraph path into the current python-docx paragraph object."""
    segments = path.split("/")
    if len(segments) < 3:
        raise ValueError(f"Invalid paragraph path '{path}'.")

    root_name = segments[0]
    if root_name == "body":
        container = document
        offset = 1
    elif root_name == "header":
        container = document.sections[int(segments[1])].header
        offset = 2
    elif root_name == "footer":
        container = document.sections[int(segments[1])].footer
        offset = 2
    else:
        raise ValueError(f"Unsupported paragraph root '{root_name}'.")

    return _resolve_in_container(container, segments[offset:])


def _iter_container_paragraphs(container, *, base_path: str, source_type: str) -> list[ParagraphTarget]:
    targets: list[ParagraphTarget] = []
    for paragraph_index, paragraph in enumerate(container.paragraphs):
        targets.append(
            ParagraphTarget(
                path=f"{base_path}/p/{paragraph_index}",
                paragraph=paragraph,
                source_type=source_type,
            )
        )
    for table_index, table in enumerate(container.tables):
        targets.extend(
            _iter_table_paragraphs(
                table,
                base_path=f"{base_path}/t/{table_index}",
                source_type=source_type,
            )
        )
    return targets


def _iter_table_paragraphs(table, *, base_path: str, source_type: str) -> list[ParagraphTarget]:
    targets: list[ParagraphTarget] = []
    header_cells = table.rows[0].cells if table.rows else []
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            header_label = None
            if row_index > 0 and len(header_cells) > cell_index:
                header_label = _join_nonempty_text(paragraph.text for paragraph in header_cells[cell_index].paragraphs)
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                targets.append(
                    ParagraphTarget(
                        path=f"{base_path}/r/{row_index}/c/{cell_index}/p/{paragraph_index}",
                        paragraph=paragraph,
                        source_type=source_type,
                        table_header_label=header_label or None,
                    )
                )
            for nested_table_index, nested_table in enumerate(cell.tables):
                targets.extend(
                    _iter_table_paragraphs(
                        nested_table,
                        base_path=f"{base_path}/r/{row_index}/c/{cell_index}/t/{nested_table_index}",
                        source_type=source_type,
                    )
                )
    return targets


def _resolve_in_container(container, segments: list[str]):
    if len(segments) < 2:
        raise ValueError("Paragraph path ended unexpectedly.")

    token = segments[0]
    index = int(segments[1])
    remaining = segments[2:]

    if token == "p":
        if remaining:
            raise ValueError("Paragraph path contains trailing tokens after paragraph selection.")
        return container.paragraphs[index]
    if token == "t":
        return _resolve_in_table(container.tables[index], remaining)
    raise ValueError(f"Unsupported paragraph path token '{token}'.")


def _resolve_in_table(table, segments: list[str]):
    if len(segments) < 4 or segments[0] != "r" or segments[2] != "c":
        raise ValueError("Table paragraph path must include row and cell indices.")

    row_index = int(segments[1])
    cell_index = int(segments[3])
    cell = table.rows[row_index].cells[cell_index]
    remaining = segments[4:]
    if len(remaining) < 2:
        raise ValueError("Table cell path ended unexpectedly.")

    token = remaining[0]
    index = int(remaining[1])
    tail = remaining[2:]
    if token == "p":
        if tail:
            raise ValueError("Paragraph path contains trailing tokens after cell paragraph selection.")
        return cell.paragraphs[index]
    if token == "t":
        return _resolve_in_table(cell.tables[index], tail)
    raise ValueError(f"Unsupported cell paragraph path token '{token}'.")


def _join_nonempty_text(parts) -> str:
    items = [str(part).strip() for part in parts if str(part).strip()]
    return " ".join(items)
