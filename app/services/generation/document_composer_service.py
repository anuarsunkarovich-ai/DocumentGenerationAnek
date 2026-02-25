"""Render resolved documents into DOCX bytes."""

import base64
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.core.exceptions import ValidationError
from app.services.generation.models import ResolvedDocument

USE_DOCUMENT_DEFAULT = object()


class DocumentComposerService:
    """Compose DOCX files from normalized constructor blocks."""

    def compose(self, document: ResolvedDocument) -> bytes:
        """Return a DOCX document for the resolved constructor payload."""
        doc = Document()
        self._apply_document_defaults(doc, document)

        for block in document.blocks:
            if block.type == "page_break":
                doc.add_page_break()
                continue

            if block.type == "header":
                self._add_header(doc, document, block.content)
            elif block.type == "text":
                self._add_text(doc, document, block.content)
            elif block.type == "table":
                self._add_table(doc, document, block.content)
            elif block.type == "image":
                self._add_image(doc, document, block.content)
            elif block.type == "signature":
                self._add_signature(doc, document, block.content)
            elif block.type == "spacer":
                self._add_spacer(doc, document, block.content)

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _apply_document_defaults(self, doc: Document, document: ResolvedDocument) -> None:
        """Apply page margins and typography defaults."""
        section = doc.sections[0]
        page = document.formatting.page
        typography = document.formatting.typography

        section.left_margin = Mm(page.margin_left_mm)
        section.right_margin = Mm(page.margin_right_mm)
        section.top_margin = Mm(page.margin_top_mm)
        section.bottom_margin = Mm(page.margin_bottom_mm)
        section.header_distance = Mm(page.header_distance_mm)
        section.footer_distance = Mm(page.footer_distance_mm)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = typography.font_family
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), typography.font_family)
        normal_style.font.size = Pt(typography.font_size_pt)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = typography.line_spacing
        normal_style.paragraph_format.first_line_indent = Mm(typography.first_line_indent_mm)
        normal_style.paragraph_format.space_before = Pt(typography.paragraph_spacing_before_pt)
        normal_style.paragraph_format.space_after = Pt(typography.paragraph_spacing_after_pt)
        normal_style.paragraph_format.alignment = self._map_alignment(typography.alignment)

    def _add_header(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add a heading paragraph."""
        paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(
            paragraph,
            document,
            alignment="center",
            first_line_indent_mm=0.0,
            keep_with_next=bool(content["style"].get("keep_with_next", True)),
        )
        run = paragraph.add_run(str(content["text"]))
        run.bold = True
        self._apply_text_style(
            paragraph,
            content["style"],
            document=document,
            default_alignment="center",
            default_first_line_indent_mm=0.0,
        )

    def _add_text(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add a text paragraph."""
        paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(paragraph, document)
        run = paragraph.add_run("" if content["text"] is None else str(content["text"]))
        style = content["style"]
        run.bold = style.get("bold", False)
        run.italic = style.get("italic", False)
        run.underline = style.get("underline", False)
        self._apply_text_style(
            paragraph,
            style,
            document=document,
            default_alignment=document.formatting.typography.alignment,
            default_first_line_indent_mm=USE_DOCUMENT_DEFAULT,
        )

    def _add_table(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add a table block."""
        columns = content["columns"]
        rows = content["rows"]
        if content.get("caption"):
            caption = doc.add_paragraph()
            self._apply_paragraph_defaults(
                caption,
                document,
                alignment="center",
                first_line_indent_mm=0.0,
                keep_with_next=True,
            )
            caption.add_run(str(content["caption"])).bold = True

        table = doc.add_table(rows=1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for index, column in enumerate(columns):
            header_cells[index].text = str(column["header"])
            paragraph = header_cells[index].paragraphs[0]
            self._apply_paragraph_defaults(
                paragraph,
                document,
                alignment="center",
                first_line_indent_mm=0.0,
            )
            for run in paragraph.runs:
                run.bold = True

        for row in rows:
            cells = table.add_row().cells
            for index, column in enumerate(columns):
                value = row.get(column["key"], "")
                cells[index].text = "" if value is None else str(value)
                paragraph = cells[index].paragraphs[0]
                self._apply_paragraph_defaults(
                    paragraph,
                    document,
                    alignment=column.get("alignment") or "left",
                    first_line_indent_mm=0.0,
                )

    def _add_image(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add an image block from a local path or base64 payload."""
        image_source = content["image"]
        if not image_source:
            raise ValidationError("Image block requires image data.")

        paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(
            paragraph,
            document,
            alignment=content["alignment"],
            first_line_indent_mm=0.0,
        )

        if not isinstance(image_source, str) or not image_source.startswith("data:image"):
            raise ValidationError("Image block requires a validated image data URL.")

        _, encoded = image_source.split(",", 1)
        image_bytes = BytesIO(base64.b64decode(encoded))
        paragraph.add_run().add_picture(image_bytes, width=Mm(content["width_mm"]))

        if content.get("caption"):
            caption = doc.add_paragraph()
            self._apply_paragraph_defaults(
                caption,
                document,
                alignment="center",
                first_line_indent_mm=0.0,
            )
            caption.add_run(str(content["caption"]))

    def _add_signature(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add a signature block."""
        alignment = (
            "right" if document.formatting.signatures_align_right else document.formatting.typography.alignment
        )

        paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(
            paragraph,
            document,
            alignment=alignment,
            first_line_indent_mm=0.0,
        )
        paragraph.add_run(f"{content['line_label']}: ____________________")

        name_paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(
            name_paragraph,
            document,
            alignment=alignment,
            first_line_indent_mm=0.0,
        )
        name_paragraph.add_run(str(content["signer_name"]))

        if content.get("signer_role"):
            role_paragraph = doc.add_paragraph()
            self._apply_paragraph_defaults(
                role_paragraph,
                document,
                alignment=alignment,
                first_line_indent_mm=0.0,
            )
            role_paragraph.add_run(str(content["signer_role"]))

        if content.get("include_date") and content.get("date"):
            date_paragraph = doc.add_paragraph()
            self._apply_paragraph_defaults(
                date_paragraph,
                document,
                alignment=alignment,
                first_line_indent_mm=0.0,
            )
            date_paragraph.add_run(str(content["date"]))

    def _add_spacer(self, doc: Document, document: ResolvedDocument, content: dict) -> None:
        """Add vertical space using paragraph spacing."""
        paragraph = doc.add_paragraph()
        self._apply_paragraph_defaults(
            paragraph,
            document,
            alignment="left",
            first_line_indent_mm=0.0,
            space_after_pt=content["height_mm"] * 2.83465,
        )

    def _apply_paragraph_defaults(
        self,
        paragraph,
        document: ResolvedDocument,
        *,
        alignment: str | None = None,
        first_line_indent_mm: float | object = USE_DOCUMENT_DEFAULT,
        keep_with_next: bool | None = None,
        space_before_pt: float | None = None,
        space_after_pt: float | None = None,
    ) -> None:
        """Apply the project GOST defaults to one paragraph."""
        typography = document.formatting.typography
        paragraph.alignment = self._map_alignment(alignment or typography.alignment)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = typography.line_spacing
        paragraph.paragraph_format.space_before = Pt(
            typography.paragraph_spacing_before_pt if space_before_pt is None else space_before_pt
        )
        paragraph.paragraph_format.space_after = Pt(
            typography.paragraph_spacing_after_pt if space_after_pt is None else space_after_pt
        )
        if first_line_indent_mm is USE_DOCUMENT_DEFAULT:
            paragraph.paragraph_format.first_line_indent = Mm(typography.first_line_indent_mm)
        else:
            paragraph.paragraph_format.first_line_indent = Mm(float(first_line_indent_mm))
        paragraph.paragraph_format.keep_with_next = bool(keep_with_next)

    def _apply_text_style(
        self,
        paragraph,
        style: dict,
        *,
        document: ResolvedDocument,
        default_alignment: str,
        default_first_line_indent_mm: float | object,
    ) -> None:
        """Apply paragraph-level style overrides."""
        resolved_alignment = style.get("alignment") or default_alignment
        resolved_first_line_indent = (
            style["first_line_indent_mm"]
            if style.get("first_line_indent_mm") is not None
            else default_first_line_indent_mm
        )
        self._apply_paragraph_defaults(
            paragraph,
            document,
            alignment=resolved_alignment,
            first_line_indent_mm=resolved_first_line_indent,
            keep_with_next=bool(style.get("keep_with_next", False)),
        )

        if style.get("font_size_pt") is not None:
            for run in paragraph.runs:
                run.font.size = Pt(style["font_size_pt"])

        if style.get("uppercase"):
            for run in paragraph.runs:
                run.text = run.text.upper()

    def _map_alignment(self, alignment: str) -> WD_ALIGN_PARAGRAPH:
        """Map alignment strings to python-docx constants."""
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
