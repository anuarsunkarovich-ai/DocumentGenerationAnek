"""Tests for strict generation formatting defaults."""

from io import BytesIO
from math import isclose

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from reportlab.lib.enums import TA_JUSTIFY, TA_RIGHT
from reportlab.lib.units import mm

from app.dtos.constructor import DocumentConstructor
from app.services.generation.document_composer_service import DocumentComposerService
from app.services.generation.models import ResolvedBlock, ResolvedDocument
from app.services.generation.pdf_render_service import (
    USE_DOCUMENT_DEFAULT,
    PdfRenderService,
)


def build_resolved_document() -> ResolvedDocument:
    """Build a small resolved document with text and signature blocks."""
    constructor = DocumentConstructor.model_validate(
        {
            "blocks": [
                {
                    "type": "text",
                    "id": "text-1",
                    "text": "A paragraph of body text.",
                },
                {
                    "type": "signature",
                    "id": "signature-1",
                    "signer_name": {"key": "signer_name", "fallback": "Anek"},
                },
            ]
        }
    )
    return ResolvedDocument(
        constructor=constructor,
        formatting=constructor.formatting,
        blocks=[
            ResolvedBlock(
                id="text-1",
                type="text",
                content={
                    "text": "A paragraph of body text.",
                    "style": {
                        "bold": False,
                        "italic": False,
                        "underline": False,
                        "uppercase": False,
                        "font_size_pt": None,
                        "alignment": None,
                        "first_line_indent_mm": None,
                        "keep_with_next": False,
                    },
                    "multiline": True,
                },
            ),
            ResolvedBlock(
                id="signature-1",
                type="signature",
                content={
                    "signer_name": "Anek",
                    "signer_role": "Director",
                    "date": "2026-03-20",
                    "line_label": "Signature",
                    "include_date": True,
                    "include_stamp_area": False,
                },
            ),
        ],
        data={},
    )


def test_docx_generation_applies_project_gost_defaults() -> None:
    """DOCX output should enforce line spacing, first-line indent, and signature alignment."""
    resolved_document = build_resolved_document()

    content = DocumentComposerService().compose(resolved_document)
    generated = DocxDocument(BytesIO(content))

    body_paragraph = generated.paragraphs[0]
    signature_line = generated.paragraphs[1]

    assert body_paragraph.text == "A paragraph of body text."
    assert body_paragraph.paragraph_format.line_spacing_rule in {
        WD_LINE_SPACING.MULTIPLE,
        WD_LINE_SPACING.ONE_POINT_FIVE,
    }
    assert body_paragraph.paragraph_format.line_spacing == 1.5
    assert isclose(body_paragraph.paragraph_format.first_line_indent.mm, 12.5, abs_tol=0.2)
    assert body_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY

    assert signature_line.text.startswith("Signature:")
    assert signature_line.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert isclose(signature_line.paragraph_format.first_line_indent.mm, 0.0, abs_tol=0.2)


def test_pdf_styles_follow_same_gost_defaults() -> None:
    """PDF style helpers should mirror the same paragraph defaults as DOCX."""
    resolved_document = build_resolved_document()
    service = PdfRenderService()
    base_style = service._build_base_style(resolved_document)

    text_style = service._build_block_style(
        resolved_document,
        base_style,
        name="BodyText",
        overrides={},
        default_alignment=resolved_document.formatting.typography.alignment,
        default_first_line_indent_mm=USE_DOCUMENT_DEFAULT,
    )
    signature_parts = service._build_signature(
        resolved_document,
        resolved_document.blocks[1].content,
        base_style,
    )

    assert text_style.alignment == TA_JUSTIFY
    assert isclose(text_style.firstLineIndent, 12.5 * mm, abs_tol=0.1)
    assert isclose(text_style.leading, 21.0, abs_tol=0.1)
    assert signature_parts[0].style.alignment == TA_RIGHT
    assert isclose(signature_parts[0].style.firstLineIndent, 0.0, abs_tol=0.1)
