"""Worker and queue orchestration tests for document generation."""

from contextlib import asynccontextmanager
from datetime import datetime, timezone
from io import BytesIO
from types import SimpleNamespace
from typing import Any, cast
from uuid import UUID, uuid4

import pytest
from fastapi import BackgroundTasks

import app.services.document_service as document_service_module
import app.services.generation.document_generation_service as generation_module
import app.workers.tasks as worker_tasks
from app.core.exceptions import ValidationError
from app.dtos.document import DocumentJobCreateRequest, ImportedTemplateDocumentJobCreateRequest
from app.dtos.template import TemplateImportBindingConfirmationItem
from app.models.enums import DocumentJobStatus
from app.services.document_service import DocumentService
from app.services.docx_template_import_service import DocxTemplateImportService
from app.services.generation.document_generation_service import (
    DocumentGenerationService,
    PermanentGenerationError,
    RetryableGenerationError,
)
from app.services.generation.models import ResolvedTemplateContext
from app.services.storage.minio import StorageError
from app.services.template_schema_service import TemplateSchemaService
from tests.conftest import build_docx_fixture
from tests.test_docx_template_import_service import _build_import_docx


def build_template_context() -> ResolvedTemplateContext:
    """Build a realistic template context for queue-orchestration tests."""
    organization_id = uuid4()
    template_id = uuid4()
    template_version_id = uuid4()
    schema = TemplateSchemaService().extract_schema(
        "certificate.docx",
        build_docx_fixture("{{student_name}}"),
    )
    return ResolvedTemplateContext(
        template_id=template_id,
        template_version_id=template_version_id,
        organization_id=organization_id,
        organization_code="math-dept",
        template_code="certificate",
        template_name="Certificate",
        template_version="1.0.0",
        original_filename="certificate.docx",
        variable_schema=schema.model_dump(mode="json"),
    )


def build_payload(context: ResolvedTemplateContext) -> DocumentJobCreateRequest:
    """Build a minimal valid job payload."""
    return DocumentJobCreateRequest.model_validate(
        {
            "organization_id": str(context.organization_id),
            "template_id": str(context.template_id),
            "template_version_id": str(context.template_version_id),
            "data": {"student_name": "Anek"},
            "constructor": {
                "blocks": [
                    {
                        "type": "text",
                        "id": "text-1",
                        "binding": {"key": "student_name"},
                    }
                ]
            },
        }
    )


class FakeTask:
    """Minimal Celery task stub for retry assertions."""

    def __init__(self, retries: int = 0) -> None:
        self.request = SimpleNamespace(retries=retries)
        self.retry_kwargs: dict[str, object] | None = None

    def retry(self, **kwargs: object) -> None:
        """Capture retry details and raise a sentinel error."""
        self.retry_kwargs = kwargs
        raise RuntimeError("retry requested")


def test_process_document_job_succeeds_without_retry() -> None:
    """Worker task should return the generation result on success."""

    class FakeService:
        async def process_job(self, job_id: UUID) -> bool:
            assert isinstance(job_id, UUID)
            return True

    result = worker_tasks._run_process_document_job(
        FakeTask(),
        job_id=uuid4(),
        service=FakeService(),
    )

    assert result is True


def test_process_document_job_retries_retryable_failure() -> None:
    """Retryable generation failures should map to Celery retry calls."""

    class FakeService:
        async def process_job(self, job_id: UUID) -> bool:
            _ = job_id
            raise RetryableGenerationError("temporary storage issue")

    task = FakeTask(retries=2)

    with pytest.raises(RuntimeError, match="retry requested"):
        worker_tasks._run_process_document_job(task, job_id=uuid4(), service=FakeService())

    assert task.retry_kwargs is not None
    assert task.retry_kwargs["countdown"] == worker_tasks.retry_delay_for_attempt(2)
    assert task.retry_kwargs["max_retries"] == worker_tasks.get_settings().worker.max_retries


def test_process_document_job_propagates_permanent_failure() -> None:
    """Permanent generation failures should not be retried by the task wrapper."""

    class FakeService:
        async def process_job(self, job_id: UUID) -> bool:
            _ = job_id
            raise PermanentGenerationError("template data is invalid")

    with pytest.raises(PermanentGenerationError, match="template data is invalid"):
        worker_tasks._run_process_document_job(FakeTask(), job_id=uuid4(), service=FakeService())


def test_recover_stale_document_jobs_requeues_recovered_ids() -> None:
    """Recovered stale jobs should be returned and re-enqueued for processing."""
    recovered_job_ids = [uuid4(), uuid4()]
    enqueued_job_ids: list[UUID] = []

    class FakeService:
        async def recover_stale_jobs(self) -> list[UUID]:
            return recovered_job_ids

    result = worker_tasks._run_recover_stale_document_jobs(
        service=FakeService(),
        enqueue_job=enqueued_job_ids.append,
    )

    assert result == [str(job_id) for job_id in recovered_job_ids]
    assert enqueued_job_ids == recovered_job_ids


def test_cleanup_maintenance_returns_summary() -> None:
    """Maintenance cleanup tasks should surface a stable summary payload."""

    class FakeService:
        async def cleanup(self):
            return SimpleNamespace(
                expired_artifacts_deleted=2,
                failed_jobs_deleted=1,
                audit_logs_deleted=3,
                temp_files_deleted=4,
                storage_bytes_reclaimed=2048,
            )

    result = worker_tasks._run_maintenance_cleanup(service=FakeService())

    assert result == {
        "expired_artifacts_deleted": 2,
        "failed_jobs_deleted": 1,
        "audit_logs_deleted": 3,
        "temp_files_deleted": 4,
        "storage_bytes_reclaimed": 2048,
    }


def test_run_billing_cycle_returns_summary() -> None:
    """Billing-cycle tasks should surface a stable summary payload."""

    class FakeService:
        async def run_billing_cycle(self):
            return SimpleNamespace(
                finalized_invoice_count=2,
                renewed_subscription_count=2,
                billed_organization_ids=[uuid4(), uuid4()],
            )

    result = worker_tasks._run_billing_cycle(service=FakeService())

    assert result["finalized_invoice_count"] == 2
    assert result["renewed_subscription_count"] == 2
    assert len(result["billed_organization_ids"]) == 2


@pytest.mark.anyio
async def test_document_generation_service_skips_duplicate_claim(monkeypatch) -> None:
    """Duplicate worker deliveries should no-op when the job cannot be claimed."""
    context = build_template_context()
    job = SimpleNamespace(
        id=uuid4(),
        organization_id=context.organization_id,
        template_id=context.template_id,
        template_version_id=context.template_version_id,
        requested_by_user_id=uuid4(),
        input_payload=build_payload(context).model_dump(mode="json"),
        status=DocumentJobStatus.PROCESSING,
        normalized_payload=None,
        cache_key=None,
        error_message=None,
        started_at=datetime.now(timezone.utc),
        completed_at=None,
        artifacts=[],
    )

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def get_by_id(self, job_id, organization_id=None):
            _ = organization_id
            if job_id == job.id:
                return job
            return None

        async def claim_for_processing(self, *, job_id, normalized_payload, cache_key, stale_before):
            _ = job_id, normalized_payload, cache_key, stale_before
            return None

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    monkeypatch.setattr(generation_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(generation_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(generation_module, "TemplateResolverService", FakeTemplateResolverService)

    service = DocumentGenerationService()
    result = await service.process_job(job.id)

    assert result is False


@pytest.mark.anyio
async def test_document_generation_service_requeues_retryable_failure(monkeypatch) -> None:
    """Transient generation errors should move the job back to queued."""
    context = build_template_context()
    job = SimpleNamespace(
        id=uuid4(),
        organization_id=context.organization_id,
        template_id=context.template_id,
        template_version_id=context.template_version_id,
        requested_by_user_id=uuid4(),
        input_payload=build_payload(context).model_dump(mode="json"),
        status=DocumentJobStatus.QUEUED,
        normalized_payload=None,
        cache_key=None,
        error_message=None,
        started_at=None,
        completed_at=None,
        artifacts=[],
    )

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def get_by_id(self, job_id, organization_id=None):
            _ = organization_id
            if job_id == job.id:
                return job
            return None

        async def claim_for_processing(self, *, job_id, normalized_payload, cache_key, stale_before):
            _ = stale_before
            if job_id != job.id:
                return None
            job.status = DocumentJobStatus.PROCESSING
            job.normalized_payload = normalized_payload
            job.cache_key = cache_key
            job.started_at = datetime.now(timezone.utc)
            return job

        async def requeue(self, job_obj, *, error_message=None):
            job_obj.status = DocumentJobStatus.QUEUED
            job_obj.error_message = error_message
            job_obj.started_at = None
            return job_obj

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    class RetryableComposer:
        def compose(self, resolved_document: object) -> bytes:
            _ = resolved_document
            raise StorageError("temporary storage issue")

    monkeypatch.setattr(generation_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(generation_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(generation_module, "TemplateResolverService", FakeTemplateResolverService)

    service = DocumentGenerationService()
    service._composer = RetryableComposer()  # type: ignore[assignment]

    with pytest.raises(RetryableGenerationError, match="temporary storage issue"):
        await service.process_job(job.id)

    assert job.status == DocumentJobStatus.QUEUED
    assert job.error_message == "temporary storage issue"
    assert job.started_at is None


@pytest.mark.anyio
async def test_document_generation_service_requeues_connection_failure(monkeypatch) -> None:
    """Connection failures should be treated as transient, such as a DB restart."""
    context = build_template_context()
    job = SimpleNamespace(
        id=uuid4(),
        organization_id=context.organization_id,
        template_id=context.template_id,
        template_version_id=context.template_version_id,
        requested_by_user_id=uuid4(),
        input_payload=build_payload(context).model_dump(mode="json"),
        status=DocumentJobStatus.QUEUED,
        normalized_payload=None,
        cache_key=None,
        error_message=None,
        started_at=None,
        completed_at=None,
        artifacts=[],
    )

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def get_by_id(self, job_id, organization_id=None):
            _ = organization_id
            if job_id == job.id:
                return job
            return None

        async def claim_for_processing(self, *, job_id, normalized_payload, cache_key, stale_before):
            _ = stale_before
            if job_id != job.id:
                return None
            job.status = DocumentJobStatus.PROCESSING
            job.normalized_payload = normalized_payload
            job.cache_key = cache_key
            job.started_at = datetime.now(timezone.utc)
            return job

        async def requeue(self, job_obj, *, error_message=None):
            job_obj.status = DocumentJobStatus.QUEUED
            job_obj.error_message = error_message
            job_obj.started_at = None
            return job_obj

    class FailingTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            _ = organization_id, template_id, template_version_id, require_published
            raise ConnectionError("database connection reset")

    monkeypatch.setattr(generation_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(generation_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(
        generation_module,
        "TemplateResolverService",
        FailingTemplateResolverService,
    )

    service = DocumentGenerationService()

    with pytest.raises(RetryableGenerationError, match="database connection reset"):
        await service.process_job(job.id)

    assert job.status == DocumentJobStatus.QUEUED
    assert job.error_message == "database connection reset"


@pytest.mark.anyio
async def test_document_generation_service_marks_permanent_failure(monkeypatch) -> None:
    """Deterministic validation errors should fail the job permanently."""
    context = build_template_context()
    job = SimpleNamespace(
        id=uuid4(),
        organization_id=context.organization_id,
        template_id=context.template_id,
        template_version_id=context.template_version_id,
        requested_by_user_id=uuid4(),
        input_payload=build_payload(context).model_dump(mode="json"),
        status=DocumentJobStatus.QUEUED,
        normalized_payload=None,
        cache_key=None,
        error_message=None,
        started_at=None,
        completed_at=None,
        artifacts=[],
    )
    audit_events: list[dict[str, object]] = []

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def get_by_id(self, job_id, organization_id=None):
            _ = organization_id
            if job_id == job.id:
                return job
            return None

        async def claim_for_processing(self, *, job_id, normalized_payload, cache_key, stale_before):
            _ = stale_before
            if job_id != job.id:
                return None
            job.status = DocumentJobStatus.PROCESSING
            job.normalized_payload = normalized_payload
            job.cache_key = cache_key
            job.started_at = datetime.now(timezone.utc)
            return job

        async def mark_failed(self, job_obj, error_message: str):
            job_obj.status = DocumentJobStatus.FAILED
            job_obj.error_message = error_message
            job_obj.completed_at = datetime.now(timezone.utc)
            return job_obj

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    class FakeAuditService:
        def __init__(self, session: object) -> None:
            _ = session

        async def log_event(self, **payload):
            audit_events.append(payload)
            return SimpleNamespace(**payload)

    class PermanentFailureComposer:
        def compose(self, resolved_document: object) -> bytes:
            _ = resolved_document
            raise ValidationError("template data is invalid")

    monkeypatch.setattr(generation_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(generation_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(generation_module, "TemplateResolverService", FakeTemplateResolverService)
    monkeypatch.setattr(generation_module, "AuditService", FakeAuditService)

    service = DocumentGenerationService()
    service._composer = PermanentFailureComposer()  # type: ignore[assignment]

    with pytest.raises(PermanentGenerationError, match="template data is invalid"):
        await service.process_job(job.id)

    assert job.status == DocumentJobStatus.FAILED
    assert job.error_message == "template data is invalid"
    assert job.completed_at is not None
    assert any(getattr(event["action"], "value", None) == "document_job_failed" for event in audit_events)


@pytest.mark.anyio
async def test_document_service_enqueues_generation_job(monkeypatch) -> None:
    """DocumentService should enqueue Celery work while preserving the API contract."""
    context = build_template_context()
    payload = build_payload(context)
    enqueued_job_ids: list[UUID] = []
    state: dict[str, object] = {}
    billing_events: list[tuple[str, UUID]] = []

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def create(self, job_obj):
            if getattr(job_obj, "id", None) is None:
                job_obj.id = uuid4()
            job_obj.created_at = datetime.now(timezone.utc)
            job_obj.artifacts = []
            state["job"] = job_obj
            return job_obj

        async def find_completed_cache_hit(self, **kwargs):
            _ = kwargs
            return None

    class FakeAuditService:
        def __init__(self, session: object) -> None:
            _ = session

        async def log_event(self, **payload):
            _ = payload
            return SimpleNamespace(**payload)

    class FakeQueueService:
        def enqueue_generation_job(
            self,
            job_id: UUID,
            *,
            organization_id=None,
            user_id=None,
            template_version_id=None,
        ) -> None:
            _ = organization_id, user_id, template_version_id
            enqueued_job_ids.append(job_id)

    class FakeBillingService:
        async def enforce_generation_allowed(self, *, organization_id, constructor, session):
            _ = constructor, session
            billing_events.append(("enforce_generation_allowed", organization_id))

        async def record_generation_request(self, *, organization_id, constructor, session):
            _ = constructor, session
            billing_events.append(("record_generation_request", organization_id))

    monkeypatch.setattr(document_service_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(document_service_module, "TemplateResolverService", FakeTemplateResolverService)
    monkeypatch.setattr(document_service_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(document_service_module, "AuditService", FakeAuditService)

    service = DocumentService()
    cast(Any, service)._job_queue_service = FakeQueueService()
    cast(Any, service)._billing_service = FakeBillingService()

    response = await service.create_job(
        payload,
        BackgroundTasks(),
        current_user_id=uuid4(),
    )

    assert response.status == DocumentJobStatus.QUEUED.value
    assert response.from_cache is False
    assert enqueued_job_ids == [response.task_id]
    assert billing_events == [
        ("enforce_generation_allowed", context.organization_id),
        ("record_generation_request", context.organization_id),
    ]


@pytest.mark.anyio
async def test_document_service_enqueues_imported_generation_job(monkeypatch) -> None:
    """Imported-template generation should queue successfully without constructor blocks."""
    context = build_template_context()
    context.render_strategy = "docx_import"
    context.storage_key = "templates/math-dept/certificate/1.0.0/certificate.docx"
    context.import_bindings = [
        {
            "binding_key": "student_name",
            "paragraph_path": "body/p/0",
            "fragment_start": 0,
            "fragment_end": 5,
            "raw_fragment": "Anek",
            "required": True,
        }
    ]
    payload = ImportedTemplateDocumentJobCreateRequest.model_validate(
        {
            "organization_id": str(context.organization_id),
            "template_id": str(context.template_id),
            "template_version_id": str(context.template_version_id),
            "data": {"student_name": "Anek"},
        }
    )
    enqueued_job_ids: list[UUID] = []
    state: dict[str, object] = {}
    billing_events: list[tuple[str, UUID, int]] = []

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def create(self, job_obj):
            if getattr(job_obj, "id", None) is None:
                job_obj.id = uuid4()
            job_obj.created_at = datetime.now(timezone.utc)
            job_obj.artifacts = []
            state["job"] = job_obj
            return job_obj

        async def find_completed_cache_hit(self, **kwargs):
            _ = kwargs
            return None

    class FakeAuditService:
        def __init__(self, session: object) -> None:
            _ = session

        async def log_event(self, **payload):
            _ = payload
            return SimpleNamespace(**payload)

    class FakeQueueService:
        def enqueue_generation_job(
            self,
            job_id: UUID,
            *,
            organization_id=None,
            user_id=None,
            template_version_id=None,
        ) -> None:
            _ = organization_id, user_id, template_version_id
            enqueued_job_ids.append(job_id)

    class FakeBillingService:
        async def enforce_generation_allowed(self, *, organization_id, constructor, session):
            _ = session
            billing_events.append(("enforce_generation_allowed", organization_id, len(constructor.blocks)))

        async def record_generation_request(self, *, organization_id, constructor, session):
            _ = session
            billing_events.append(("record_generation_request", organization_id, len(constructor.blocks)))

    monkeypatch.setattr(document_service_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(document_service_module, "TemplateResolverService", FakeTemplateResolverService)
    monkeypatch.setattr(document_service_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(document_service_module, "AuditService", FakeAuditService)

    service = DocumentService()
    cast(Any, service)._job_queue_service = FakeQueueService()
    cast(Any, service)._billing_service = FakeBillingService()

    response = await service.create_imported_job(
        payload,
        BackgroundTasks(),
        current_user_id=uuid4(),
    )

    assert response.status == DocumentJobStatus.QUEUED.value
    assert response.from_cache is False
    assert enqueued_job_ids == [response.task_id]
    assert billing_events == [
        ("enforce_generation_allowed", context.organization_id, 0),
        ("record_generation_request", context.organization_id, 0),
    ]


@pytest.mark.anyio
async def test_document_generation_service_processes_imported_docx_job(monkeypatch) -> None:
    """Imported-DOCX jobs should render from the stored source file and skip PDF output."""
    source_bytes = _build_import_docx()
    import_service = DocxTemplateImportService()
    analysis = import_service.analyze("invoice.docx", source_bytes)
    bindings, schema = import_service.confirm_bindings(
        analysis=analysis,
        confirmations=[
            TemplateImportBindingConfirmationItem(
                candidate_id=candidate.id,
                binding_key=candidate.suggested_binding,
            )
            for candidate in analysis.candidates
        ],
    )
    context = ResolvedTemplateContext(
        template_id=uuid4(),
        template_version_id=uuid4(),
        organization_id=uuid4(),
        organization_code="finance",
        template_code="invoice",
        template_name="Invoice",
        template_version="1.0.0",
        original_filename="invoice.docx",
        variable_schema=schema.model_dump(mode="json"),
        storage_key="templates/finance/invoice/1.0.0/invoice.docx",
        render_strategy="docx_import",
        import_bindings=[binding.model_dump(mode="json") for binding in bindings],
    )
    job = SimpleNamespace(
        id=uuid4(),
        organization_id=context.organization_id,
        template_id=context.template_id,
        template_version_id=context.template_version_id,
        requested_by_user_id=uuid4(),
        input_payload={
            "generation_mode": "docx_import",
            "data": {
                "client_name": "Anek LLC",
                "invoice_date": "2026-03-21",
                "amount": "1500 USD",
            },
        },
        status=DocumentJobStatus.QUEUED,
        normalized_payload=None,
        cache_key=None,
        error_message=None,
        started_at=None,
        completed_at=None,
        artifacts=[],
    )
    audit_events: list[dict[str, object]] = []
    rendered_docx: dict[str, bytes] = {}

    @asynccontextmanager
    async def fake_transaction_session():
        yield object()

    class FakeDocumentRepository:
        def __init__(self, session: object) -> None:
            _ = session

        async def get_by_id(self, job_id, organization_id=None):
            _ = organization_id
            if job_id == job.id:
                return job
            return None

        async def claim_for_processing(self, *, job_id, normalized_payload, cache_key, stale_before):
            _ = stale_before
            if job_id != job.id:
                return None
            job.status = DocumentJobStatus.PROCESSING
            job.normalized_payload = normalized_payload
            job.cache_key = cache_key
            job.started_at = datetime.now(timezone.utc)
            return job

        async def mark_completed(self, job_obj):
            job_obj.status = DocumentJobStatus.COMPLETED
            job_obj.completed_at = datetime.now(timezone.utc)
            return job_obj

    class FakeTemplateResolverService:
        def __init__(self, session: object) -> None:
            _ = session

        async def resolve(
            self,
            *,
            organization_id,
            template_id,
            template_version_id,
            require_published: bool = False,
        ):
            assert require_published is False
            assert organization_id == context.organization_id
            assert template_id == context.template_id
            assert template_version_id == context.template_version_id
            return context

    class FakeStorageService:
        async def download_bytes(self, key: str) -> bytes:
            assert key == context.storage_key
            return source_bytes

    class FakeArtifactService:
        def __init__(self, session: object, storage_service: FakeStorageService) -> None:
            _ = session, storage_service

        async def store_docx(self, *, context, job_id, user_id, content):
            _ = context, job_id, user_id
            rendered_docx["content"] = content
            return SimpleNamespace(kind="docx")

        async def store_pdf(self, *, context, job_id, user_id, content):
            _ = context, job_id, user_id, content
            raise AssertionError("Imported DOCX generation should not store a PDF artifact.")

    class FakeAuditService:
        def __init__(self, session: object) -> None:
            _ = session

        async def log_event(self, **payload):
            audit_events.append(payload)
            return SimpleNamespace(**payload)

    class FakeBillingService:
        async def enforce_storage_delta_allowed(self, *, organization_id, additional_bytes, session):
            _ = organization_id, additional_bytes, session
            return None

    monkeypatch.setattr(generation_module, "get_transaction_session", fake_transaction_session)
    monkeypatch.setattr(generation_module, "DocumentRepository", FakeDocumentRepository)
    monkeypatch.setattr(generation_module, "TemplateResolverService", FakeTemplateResolverService)
    monkeypatch.setattr(generation_module, "ArtifactService", FakeArtifactService)
    monkeypatch.setattr(generation_module, "AuditService", FakeAuditService)
    monkeypatch.setattr(generation_module, "BillingService", FakeBillingService)

    service = DocumentGenerationService()
    service._storage_service = FakeStorageService()  # type: ignore[assignment]

    result = await service.process_job(job.id)

    assert result is True
    assert job.status == DocumentJobStatus.COMPLETED
    assert "content" in rendered_docx
    from docx import Document

    rendered = Document(BytesIO(rendered_docx["content"]))
    assert rendered.paragraphs[0].text == "Client Name: Anek LLC"
    assert rendered.paragraphs[1].text == "Invoice Date: 2026-03-21"
    assert rendered.tables[0].rows[1].cells[1].text == "1500 USD"
    assert any(
        event["action"] == generation_module.AuditAction.DOCUMENT_JOB_COMPLETED
        and event["payload"]["generation_mode"] == "docx_import"
        for event in audit_events
    )
